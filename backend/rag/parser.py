from pydantic import BaseModel
from typing import List
import fitz # PyMuPDF
import docx
import io

class ParsedPage(BaseModel):
    page_number: int
    text: str

class ParsedDocument(BaseModel):
    text: str
    pages: List[ParsedPage]
    mime_type: str

def parse_pdf(file_bytes: bytes) -> ParsedDocument:
    doc = fitz.open("pdf", file_bytes)
    pages = []
    full_text = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages.append(ParsedPage(page_number=page_num + 1, text=text))
            full_text.append(text)
            
    return ParsedDocument(
        text="\n\n".join(full_text),
        pages=pages,
        mime_type="application/pdf"
    )

def parse_docx(file_bytes: bytes) -> ParsedDocument:
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # docx doesn't have a strict concept of pages unless calculated,
    # so we treat the whole document as page 1.
    pages = [ParsedPage(page_number=1, text=full_text)]
    
    return ParsedDocument(
        text=full_text,
        pages=pages,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def parse_text(file_bytes: bytes, mime_type: str) -> ParsedDocument:
    text = file_bytes.decode('utf-8', errors='ignore')
    pages = [ParsedPage(page_number=1, text=text)]
    return ParsedDocument(
        text=text,
        pages=pages,
        mime_type=mime_type
    )

def parse_document(file_bytes: bytes, mime_type: str) -> ParsedDocument:
    """
    Dispatcher for parsing based on MIME type.
    """
    if mime_type == "application/pdf":
        return parse_pdf(file_bytes)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(file_bytes)
    elif mime_type in ["text/plain", "text/markdown"]:
        return parse_text(file_bytes, mime_type)
    else:
        raise ValueError(f"Unsupported MIME type for parsing: {mime_type}")
